"""
record_generator.py
Generates professional .docx service documents for:
  - Load Test records
  - Repair / Inspection records
  - Inventory records
Matching the Excelsior Aviation GSE Corporation brand style.
"""
import io
import os
import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIR  = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(BASE_DIR, 'logo.jpeg')

BLUE  = RGBColor(0x00, 0x33, 0xCC)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY  = RGBColor(0x44, 0x44, 0x44)
ORANGE = RGBColor(0xFF, 0x6B, 0x00)


# ── Helpers ────────────────────────────────────────────────────────────────

def _setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(11)
    return doc


def _para(doc, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment         = align
    pf.space_before      = Pt(space_before)
    pf.space_after       = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def _run(para, text, bold=False, size=11, color=BLACK, underline=False, italic=False):
    r = para.add_run(str(text))
    r.bold         = bold
    r.underline    = underline
    r.italic       = italic
    r.font.name    = 'Times New Roman'
    r.font.size    = Pt(size)
    r.font.color.rgb = color
    return r


def _header(doc):
    """Company letterhead."""
    if os.path.exists(LOGO_PATH):
        p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
        p.add_run().add_picture(LOGO_PATH, height=Inches(0.85))

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _run(p, 'Excelsior Aviation GSE Corporation', bold=True, size=22, color=BLUE)

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _run(p, 'Lot \u2013 9 Dona Lucing Subdivision Barangay Calibutbut, Bacolor 2001 Pampanga', size=9, color=BLUE)

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _run(p, 'Email Address: sale.service@excelsioraviation.com', size=9, color=BLUE)

    # Horizontal rule
    p = _para(doc, space_after=8)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0033CC')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _section_title(doc, title, icon=''):
    p = _para(doc, space_before=6, space_after=4)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'FF6B00')
    pPr.append(shd)
    _run(p, f'  {title}  ', bold=True, size=11, color=WHITE)


def _info_row(table_rows, label, value, bold_val=False):
    """Add a label/value row to a 2-col info table."""
    row = table_rows.add_row()
    # Label cell
    lc = row.cells[0]
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    _run(lp, label, bold=True, size=10, color=GRAY)
    # Value cell
    vc = row.cells[1]
    vp = vc.paragraphs[0]
    vp.paragraph_format.space_after = Pt(2)
    _run(vp, str(value) if value else '—', bold=bold_val, size=10)


def _info_table(doc, rows_data):
    """Create a clean 2-column info table: [(label, value), ...]"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Inches(2.2)
    tbl.columns[1].width = Inches(4.3)
    for label, value in rows_data:
        row = tbl.add_row()
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)
        # Shade label cell
        tc = row.cells[0]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'FFF3EC')
        tcPr.append(shd)
        lp = row.cells[0].paragraphs[0]
        lp.paragraph_format.space_after = Pt(2)
        _run(lp, str(label), bold=True, size=10, color=GRAY)
        vp = row.cells[1].paragraphs[0]
        vp.paragraph_format.space_after = Pt(2)
        _run(vp, str(value) if value else '\u2014', size=10)
    return tbl


def _fmt_date(d):
    if not d:
        return ''
    try:
        if isinstance(d, str):
            d = datetime.date.fromisoformat(d)
        return d.strftime('%B %d, %Y')
    except Exception:
        return str(d)


def _footer_line(doc):
    _para(doc, space_before=8, space_after=2)
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _run(p, 'Excelsior Aviation GSE Corporation \u2014 Authorized Service Record',
         italic=True, size=8, color=GRAY)


def _save(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════════════════════
# 1. LOAD TEST RECORD
# ══════════════════════════════════════════════════════════════════════════

def generate_loadtest_doc(record):
    doc = _setup_doc()
    _header(doc)

    # Title
    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
    _run(p, 'LOAD TEST SERVICE RECORD', bold=True, size=16)

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _run(p, f'Certificate No.: ', bold=False, size=11)
    _run(p, str(record.certificate_number or '\u2014'), bold=True, size=11, color=ORANGE)

    # Customer Info
    _section_title(doc, 'CUSTOMER INFORMATION')
    _info_table(doc, [
        ('Customer Name',    record.customer_name),
        ('Customer Address', record.customer_address),
    ])
    _para(doc, space_after=6)

    # Equipment Details
    _section_title(doc, 'EQUIPMENT DETAILS')
    _info_table(doc, [
        ('Company',          record.company),
        ('Description',      record.description),
        ('Model Number',     record.model_number),
        ('Serial Number',    record.serial_number),
        ('Equipment',        record.equipment),
    ])
    _para(doc, space_after=6)

    # Load Test Results
    _section_title(doc, 'LOAD TEST RESULTS')
    tbl = doc.add_table(rows=2, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    headers = ['TESTED LOAD', 'SAFE WORKING LOAD', 'DATE INSPECTION']
    values  = [record.tested_load, record.safe_load, _fmt_date(record.date_inspection)]
    for i, (h, v) in enumerate(zip(headers, values)):
        hc = tbl.rows[0].cells[i]
        tc = hc._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'FF6B00')
        tcPr.append(shd)
        hp = hc.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(hp, h, bold=True, size=9, color=WHITE)
        vc = tbl.rows[1].cells[i]
        vp = vc.paragraphs[0]; vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(vp, str(v) if v else '\u2014', bold=True, size=11)
    _para(doc, space_after=6)

    # Inspection
    _section_title(doc, 'INSPECTION INFO')
    _info_table(doc, [
        ('Mechanic / Technician', record.mechanic),
        ('Date Inspection',       _fmt_date(record.date_inspection)),
        ('Date Due',              _fmt_date(record.date_due)),
        ('Certificate No.',       record.certificate_number),
    ])
    _para(doc, space_after=6)

    # Remarks
    if record.remark:
        _section_title(doc, 'REMARKS')
        p = _para(doc, space_after=6)
        p.paragraph_format.left_indent = Inches(0.2)
        _run(p, record.remark, size=10)

    # Signature
    _para(doc, space_before=12, space_after=2)
    p = _para(doc)
    _run(p, 'Prepared by: ', bold=True, size=10)
    if record.created_by:
        name = record.created_by.get_full_name() or record.created_by.username
        _run(p, name, size=10, underline=True)
        _run(p, f'    Date Generated: {_fmt_date(record.created_at)}', size=10)

    _footer_line(doc)
    return _save(doc)


# ══════════════════════════════════════════════════════════════════════════
# 2. REPAIR / INSPECTION RECORD
# ══════════════════════════════════════════════════════════════════════════

def generate_repair_doc(record):
    doc = _setup_doc()
    _header(doc)

    record_type_label = 'REPAIR' if record.record_type == 'repair' else 'INSPECTION'

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
    _run(p, f'{record_type_label} SERVICE RECORD', bold=True, size=16)

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    _run(p, 'Report No.: ', bold=False, size=11)
    _run(p, str(record.report_number or '\u2014'), bold=True, size=11, color=ORANGE)

    # Customer
    _section_title(doc, 'CUSTOMER INFORMATION')
    _info_table(doc, [
        ('Customer Name',    record.customer_name),
        ('Customer Address', record.customer_address),
    ])
    _para(doc, space_after=6)

    # Equipment
    _section_title(doc, 'EQUIPMENT DETAILS')
    _info_table(doc, [
        ('Company',      record.company),
        ('Description',  record.description),
        ('Model Number', record.model_number),
        ('Serial Number',record.serial_number),
        ('Report No.',   record.report_number),
        ('Mechanic',     record.mechanic),
        ('Date',         _fmt_date(record.date)),
    ])
    _para(doc, space_after=6)

    # Customer Report
    if record.customer_report:
        _section_title(doc, 'CUSTOMER REPORT / COMPLAINT')
        p = _para(doc, space_after=6)
        p.paragraph_format.left_indent = Inches(0.2)
        _run(p, record.customer_report, size=10)

    # Diagnose
    if record.diagnose_result:
        _section_title(doc, 'DIAGNOSIS / FINDINGS')
        p = _para(doc, space_after=6)
        p.paragraph_format.left_indent = Inches(0.2)
        _run(p, record.diagnose_result, size=10)

    # Remarks
    if record.remarks:
        _section_title(doc, 'REMARKS')
        p = _para(doc, space_after=6)
        p.paragraph_format.left_indent = Inches(0.2)
        _run(p, record.remarks, size=10)

    # Signature
    _para(doc, space_before=12, space_after=2)
    p = _para(doc)
    _run(p, 'Prepared by: ', bold=True, size=10)
    if record.created_by:
        name = record.created_by.get_full_name() or record.created_by.username
        _run(p, name, size=10, underline=True)
        _run(p, f'    Date Generated: {_fmt_date(record.created_at)}', size=10)

    _footer_line(doc)
    return _save(doc)


# ══════════════════════════════════════════════════════════════════════════
# 3. INVENTORY RECORD
# ══════════════════════════════════════════════════════════════════════════

def generate_inventory_doc(record):
    doc = _setup_doc()
    _header(doc)

    p = _para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=10)
    _run(p, 'INVENTORY RECORD', bold=True, size=16)

    # Customer
    _section_title(doc, 'CUSTOMER INFORMATION')
    _info_table(doc, [
        ('Customer Name',    record.customer_name),
        ('Customer Address', record.customer_address),
    ])
    _para(doc, space_after=6)

    # Item Details
    _section_title(doc, 'ITEM DETAILS')
    _info_table(doc, [
        ('Description',  record.description),
        ('Model Number', record.model_number),
        ('Serial Number',record.serial_number),
        ('Location',     record.location),
        ('Quantity',     record.quantity),
    ])
    _para(doc, space_after=6)

    # Remarks
    if record.remarks:
        _section_title(doc, 'REMARKS')
        p = _para(doc, space_after=6)
        p.paragraph_format.left_indent = Inches(0.2)
        _run(p, record.remarks, size=10)

    # Signature
    _para(doc, space_before=12, space_after=2)
    p = _para(doc)
    _run(p, 'Prepared by: ', bold=True, size=10)
    if record.created_by:
        name = record.created_by.get_full_name() or record.created_by.username
        _run(p, name, size=10, underline=True)
        _run(p, f'    Date Generated: {_fmt_date(record.created_at)}', size=10)

    _footer_line(doc)
    return _save(doc)