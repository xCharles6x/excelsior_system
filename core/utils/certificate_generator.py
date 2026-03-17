"""
core/utils/certificate_generator.py

Generates a certificate .docx by copying the original template and
filling in all database values using python-docx.

Place the template at:  core/utils/certificate_template.docx

Requirements:
    pip install python-docx
"""

import io
import os
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'certificate_template.docx')


def _fmt_date(d):
    """Format date → 'January 06, 2026'"""
    if not d:
        return ''
    if isinstance(d, str):
        return d
    return d.strftime('%B %d, %Y')


def _s(value):
    """Safe string — returns '' for None."""
    return str(value).strip() if value is not None else ''


def _set_run_text(run, text):
    """Set a run's text without disturbing its formatting."""
    run.text = text


def _clear_runs_after(para, keep_runs):
    """Remove all runs after index `keep_runs` in the paragraph."""
    runs = para.runs
    for run in runs[keep_runs:]:
        run._element.getparent().remove(run._element)


def _add_value_run(para, text, copy_from_run=None):
    """
    Add a new run with `text` at the end of `para`,
    copying formatting from `copy_from_run` if provided.
    """
    new_run = para.add_run(text)
    if copy_from_run:
        rpr = copy_from_run._r.get_or_add_rPr()
        new_rpr = copy.deepcopy(rpr)
        new_run._r.get_or_add_rPr().getparent().replace(
            new_run._r.get_or_add_rPr(), new_rpr
        )
    return new_run


def _fill_para_after_colon(para, value):
    """
    For label paragraphs like 'CUSTOMER : [tabs] ADDRESS : [tabs]',
    inject value text into the run right after the first colon run.
    Strategy: find the run containing ': ' and insert value after it,
    clearing any subsequent tab-only runs up to the next label.
    """
    if not value:
        return
    runs = para.runs
    colon_idx = None
    for i, run in enumerate(runs):
        if run.text.strip() in (':', ': '):
            colon_idx = i
            break
    if colon_idx is None:
        # Try finding label ending with colon
        for i, run in enumerate(runs):
            if run.text.rstrip().endswith(':'):
                colon_idx = i
                break
    if colon_idx is not None and colon_idx + 1 < len(runs):
        # Replace the first tab run after colon with the value
        next_run = runs[colon_idx + 1]
        if next_run.text.strip() == '' or next_run.text == '\t':
            next_run.text = ' ' + value


def generate_certificate(cert, rows):
    """
    Fill the certificate template with data from cert and rows.
    Returns a BytesIO buffer of the completed .docx.
    """
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"Certificate template not found at {TEMPLATE_PATH}. "
            "Copy your .docx template to core/utils/certificate_template.docx"
        )

    doc = Document(TEMPLATE_PATH)
    paras = doc.paragraphs

    # ── 1. Certificate Number (para 8) ───────────────────────────────────────
    # Runs: [0]='CERTIFICATE NO. ' [1-6]=red bold number parts
    p8 = paras[8]
    runs = p8.runs
    # Clear runs 1 onwards (the old number) and write new cert number
    cert_num = _s(cert.certificate_number)
    if len(runs) >= 2:
        # Keep run[0] as label, set run[1] to full cert number, remove rest
        runs[1].text = cert_num
        for run in runs[2:]:
            run.text = ''

    # ── 2. Customer & Address (para 11) ──────────────────────────────────────
    # Structure: CUSTOMER [tab][tab]: [tab*7] ADDRESS [tab][tab]: [tab]
    # run indices: 0=CUSTOMER  1,2=tabs  3=': '  4..10=tabs  11=ADDRESS  12,13=tabs  14=': '  15=tab
    p11 = paras[11]
    runs = p11.runs
    customer_name = _s(cert.customer.name if cert.customer_id else '')
    address_val   = _s(cert.address)

    # Inject customer name: replace run[4] (first tab after ':')
    if len(runs) > 4:
        runs[4].text = ' ' + customer_name
        # Clear remaining tab runs up to ADDRESS label (runs 5-10)
        for i in range(5, min(11, len(runs))):
            if runs[i].text == '\t':
                runs[i].text = ''

    # Inject address: replace run[15] (tab after ADDRESS ':')
    if len(runs) > 15:
        runs[15].text = ' ' + address_val

    # ── 3. Date of Inspection (para 13) ──────────────────────────────────────
    # 'DATE OF INSPECTION \t:\t'  → inject after ':'
    p13 = paras[13]
    runs13 = p13.runs
    doi = _fmt_date(cert.date_of_inspection)
    # Find the ':' run and put value in the next run
    for i, run in enumerate(runs13):
        if ':' in run.text and i + 1 < len(runs13):
            runs13[i + 1].text = ' ' + doi
            break
        elif run.text.rstrip().endswith(':'):
            run.text = run.text + ' ' + doi
            break

    # ── 4. Due Next Inspection (para 14) ─────────────────────────────────────
    p14 = paras[14]
    runs14 = p14.runs
    dni = _fmt_date(cert.due_next_inspection)
    for i, run in enumerate(runs14):
        if ':' in run.text:
            if i + 1 < len(runs14):
                runs14[i + 1].text = ' ' + dni
            else:
                run.text = run.text.rstrip() + ' ' + dni
            break

    # ── 5. Product Description (para 15) ─────────────────────────────────────
    p15 = paras[15]
    runs15 = p15.runs
    for i, run in enumerate(runs15):
        if ':' in run.text:
            if i + 1 < len(runs15):
                runs15[i + 1].text = ' ' + _s(cert.product_description)
            else:
                run.text = run.text.rstrip() + ' ' + _s(cert.product_description)
            break

    # ── 6. Capacity (para 16) ────────────────────────────────────────────────
    p16 = paras[16]
    runs16 = p16.runs
    for i, run in enumerate(runs16):
        if ':' in run.text:
            if i + 1 < len(runs16):
                runs16[i + 1].text = ' ' + _s(cert.capacity)
            else:
                run.text = run.text.rstrip() + ' ' + _s(cert.capacity)
            break

    # ── 7. Model Number (para 17) ────────────────────────────────────────────
    p17 = paras[17]
    runs17 = p17.runs
    for i, run in enumerate(runs17):
        if ':' in run.text:
            if i + 1 < len(runs17):
                runs17[i + 1].text = ' ' + _s(cert.model_number)
            else:
                run.text = run.text.rstrip() + ' ' + _s(cert.model_number)
            break

    # ── 8. Serial Number (para 18) ───────────────────────────────────────────
    p18 = paras[18]
    runs18 = p18.runs
    for i, run in enumerate(runs18):
        if ':' in run.text:
            if i + 1 < len(runs18):
                runs18[i + 1].text = ' ' + _s(cert.serial_number)
            else:
                run.text = run.text.rstrip() + ' ' + _s(cert.serial_number)
            break

    # ── 9. TEST RESULTS table (Table 0) ──────────────────────────────────────
    # row 2 → [WORKING LOAD LIMITS value, TESTED LOAD value, PRESSURE RELIEF value]
    t0 = doc.tables[0]
    if len(t0.rows) > 2:
        data_row = t0.rows[2]
        _set_cell_value(data_row.cells[0], _s(cert.working_load_limits))
        _set_cell_value(data_row.cells[1], _s(cert.tested_load))
        _set_cell_value(data_row.cells[2], _s(cert.pressure_relief))

    # ── 10. TEST EQUIPMENT table (Table 1) ───────────────────────────────────
    # row 2 is the first data row (row 0 = title, row 1 = headers)
    t1 = doc.tables[1]
    rows_list = list(rows)
    for i, eq_row in enumerate(rows_list):
        # Ensure enough rows exist (template has 1 data row; add more if needed)
        if i + 2 >= len(t1.rows):
            _add_table_row(t1, i + 2)
        data_row = t1.rows[i + 2]
        _set_cell_value(data_row.cells[0], _s(eq_row.description))
        _set_cell_value(data_row.cells[1], _s(eq_row.capacity))
        _set_cell_value(data_row.cells[2], _s(eq_row.model_no))
        _set_cell_value(data_row.cells[3], _s(eq_row.serial_no))
        _set_cell_value(data_row.cells[4], _s(eq_row.part_no))
        _set_cell_value(data_row.cells[5], _s(eq_row.expiry_date))

    # If no equipment rows, clear the placeholder text in row 2
    if not rows_list and len(t1.rows) > 2:
        for cell in t1.rows[2].cells:
            _set_cell_value(cell, '')

    # ── 11. Load Cell Certificate (Table 2) ──────────────────────────────────
    t2 = doc.tables[2]
    if len(t2.rows) > 0 and len(t2.rows[0].cells) > 1:
        _set_cell_value(t2.rows[0].cells[1], _s(cert.load_cell_certificate_number))

    # ── 12. Work Perform (paras 25-30 area) ──────────────────────────────────
    # Para 25 starts the work perform text; replace with cert.work_perform
    if cert.work_perform:
        wp_lines = cert.work_perform.strip().split('\n')
        # Para 25 = first line of work perform text
        if len(paras) > 25:
            _replace_para_text(paras[25], wp_lines[0] if wp_lines else '')

    # ── 13. Technician Name & Date (para 34) ─────────────────────────────────
    # run[0] = 'Ariel Lacsina'  run[10-14] = 'January 06, 2026' parts
    p34 = paras[34]
    runs34 = p34.runs
    if runs34:
        runs34[0].text = _s(cert.technician_name)
    # Replace the date runs (10-14) with single run
    tech_date = _fmt_date(cert.technician_date)
    if len(runs34) > 10:
        runs34[10].text = tech_date
        for run in runs34[11:]:
            run.text = ''

    # ── Save to buffer ────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_cell_value(cell, text):
    """Set the first paragraph of a table cell to `text`, preserving formatting."""
    para = cell.paragraphs[0]
    runs = para.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ''
    else:
        run = para.add_run(text)
        # Apply basic Times New Roman 10pt formatting to match template
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)


def _replace_para_text(para, text):
    """Replace all run text in a paragraph with `text` in the first run."""
    runs = para.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ''
    else:
        para.add_run(text)


def _add_table_row(table, copy_from_row_idx=2):
    """Copy a row from the table and append it as a new row."""
    try:
        source_row = table.rows[copy_from_row_idx]
        new_tr = copy.deepcopy(source_row._tr)
        table._tbl.append(new_tr)
        # Clear any text in the new row
        new_row = table.rows[-1]
        for cell in new_row.cells:
            _set_cell_value(cell, '')
    except Exception:
        pass